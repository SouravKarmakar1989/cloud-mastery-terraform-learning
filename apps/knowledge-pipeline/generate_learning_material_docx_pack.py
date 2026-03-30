"""
Generate blog-style DOCX learning material from KB markdown files.

Focus:
- Preserve coverage of all topics from markdown files.
- Add cross-cloud teaching (AWS, Azure, GCP) throughout.
- Produce a hierarchical learning pack under Learning Material/.
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document


SERVICE_EQUIVALENTS = {
    "Amazon EC2": ("Azure Virtual Machines", "Compute Engine"),
    "AWS Lambda": ("Azure Functions", "Cloud Functions / Cloud Run"),
    "Amazon S3": ("Azure Blob Storage", "Cloud Storage"),
    "Amazon RDS": ("Azure SQL / PostgreSQL", "Cloud SQL"),
    "Amazon DynamoDB": ("Cosmos DB", "Firestore / Bigtable"),
    "Amazon EKS": ("AKS", "GKE"),
    "Amazon ECS": ("Azure Container Apps / AKS", "Cloud Run / GKE"),
    "Elastic Load Balancing": ("Azure Load Balancer / Application Gateway", "Cloud Load Balancing"),
    "Amazon CloudWatch": ("Azure Monitor", "Cloud Monitoring"),
    "AWS CloudFormation": ("ARM/Bicep", "Deployment Manager / IaC tools"),
    "AWS CDK": ("Bicep/Terraform-based app modeling", "Terraform/Pulumi patterns"),
    "AWS SAM": ("Azure Functions templates", "Cloud Run/Functions templates"),
    "Amazon VPC": ("Azure Virtual Network", "VPC Network"),
    "Amazon EBS": ("Managed Disks", "Persistent Disk"),
    "Amazon EFS": ("Azure Files / NetApp Files", "Filestore"),
}


@dataclass
class ModuleBlock:
    name: str
    source_files: list[str] = field(default_factory=list)
    transcript_notes: list[str] = field(default_factory=list)
    polished_notes: list[str] = field(default_factory=list)
    synthesis: list[str] = field(default_factory=list)


@dataclass
class ParsedKB:
    title: str
    topic: str
    source_md: Path
    modules: list[ModuleBlock] = field(default_factory=list)
    concepts: list[str] = field(default_factory=list)
    services: list[str] = field(default_factory=list)
    referrals: list[str] = field(default_factory=list)
    key_insights: list[str] = field(default_factory=list)
    hands_on_activities: list[dict[str, Any]] = field(default_factory=list)
    cross_cloud_rows: list[tuple[str, str, str]] = field(default_factory=list)


def clean_line(line: str) -> str:
    return re.sub(r"\s+", " ", line.strip())


def detect_service_in_text(text: str) -> list[str]:
    found: list[str] = []
    lower = text.lower()
    patterns = {
        "Amazon EC2": ["ec2", "instance"],
        "AWS Lambda": ["lambda", "serverless"],
        "Amazon S3": ["s3", "bucket", "object storage"],
        "Amazon RDS": ["rds", "aurora", "relational"],
        "Amazon DynamoDB": ["dynamodb", "nosql"],
        "Amazon EKS": ["eks", "kubernetes"],
        "Amazon ECS": ["ecs", "fargate"],
        "Elastic Load Balancing": ["load balancer", "alb", "nlb", "elb"],
        "Amazon CloudWatch": ["cloudwatch", "metrics", "alarm", "logs"],
        "AWS CloudFormation": ["cloudformation", "stack"],
        "AWS CDK": ["cdk"],
        "AWS SAM": ["sam"],
        "Amazon VPC": ["vpc", "subnet", "route table"],
        "Amazon EBS": ["ebs", "volume"],
        "Amazon EFS": ["efs", "file system"],
    }
    for service, keys in patterns.items():
        if any(k in lower for k in keys):
            found.append(service)
    return found


def parse_md(path: Path) -> ParsedKB:
    lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()

    title = path.stem
    topic = path.stem.replace("_", " ")
    for line in lines:
        if line.startswith("# "):
            title = clean_line(line[2:])
            break
    for line in lines:
        if line.startswith("- Topic:"):
            topic = clean_line(line.split(":", 1)[1])
            break

    parsed = ParsedKB(title=title, topic=topic, source_md=path)

    current_module: ModuleBlock | None = None
    section = ""
    in_services = False
    in_concepts = False
    in_referrals = False

    current_activity: dict[str, Any] | None = None

    for raw in lines:
        line = clean_line(raw)
        if not line:
            continue

        if line.startswith("#### Module:"):
            module_name = clean_line(line.split(":", 1)[1])
            current_module = ModuleBlock(name=module_name)
            parsed.modules.append(current_module)
            section = ""
            continue

        if line.startswith("##### Source Transcript Details"):
            section = "source"
            continue
        if line.startswith("##### Transcript-Enriched Learning Notes"):
            section = "transcript_notes"
            continue
        if line.startswith("##### Polished Architect Notes"):
            section = "polished_notes"
            continue
        if line.startswith("##### Architect Synthesis"):
            section = "synthesis"
            continue

        if line.startswith("### 1. Concepts List"):
            in_concepts = True
            in_services = False
            continue
        if line.startswith("### 2. Services List"):
            in_services = True
            in_concepts = False
            continue
        if line.startswith("### 3."):
            in_services = False
            in_concepts = False

        if line.startswith("## Cross-Domain Referrals"):
            in_referrals = True
            continue
        if line.startswith("## Step 1 - Line-by-Line Extraction"):
            in_referrals = False

        if line.startswith("### Activity "):
            current_activity = {
                "title": line.replace("### ", ""),
                "meta": [],
                "steps": [],
            }
            parsed.hands_on_activities.append(current_activity)
            continue

        if current_activity is not None and line.startswith("### ") and not line.startswith("### Activity "):
            current_activity = None

        if current_activity is not None:
            if line.startswith("- "):
                current_activity["meta"].append(line[2:])
                continue
            if re.match(r"^\d+\.\s", line):
                current_activity["steps"].append(re.sub(r"^\d+\.\s", "", line))
                continue

        if in_concepts and line.startswith("- ") and "(No explicit" not in line:
            parsed.concepts.append(line[2:])
            continue

        if in_services and line.startswith("- ") and "(No explicit" not in line:
            parsed.services.append(line[2:])
            continue

        if in_referrals and line.startswith("- outputs/"):
            parsed.referrals.append(line[2:])
            continue

        if line.startswith("- Key Insights:"):
            parsed.key_insights.append(clean_line(line.split(":", 1)[1]))
            continue

        if line.startswith("| ") and "|" in line and not line.startswith("|---"):
            parts = [p.strip() for p in line.strip("|").split("|")]
            if len(parts) == 4 and parts[0] != "Capability":
                parsed.cross_cloud_rows.append((parts[0], parts[1], parts[2] + " / " + parts[3]))

        if current_module is None:
            continue

        if section == "source" and line.startswith("- outputs/"):
            current_module.source_files.append(line[2:])
        elif section == "transcript_notes" and line.startswith("- "):
            current_module.transcript_notes.append(line[2:])
        elif section == "polished_notes" and line.startswith("- "):
            current_module.polished_notes.append(line[2:])
        elif section == "synthesis" and line.startswith("- "):
            current_module.synthesis.append(line[2:])

    if not parsed.services:
        service_guess: list[str] = []
        for m in parsed.modules:
            for note in m.transcript_notes[:80]:
                service_guess.extend(detect_service_in_text(note))
        parsed.services = sorted(set(service_guess))

    return parsed


def add_bullet(doc: Any, text: str) -> None:
    doc.add_paragraph(f"- {text}")


def add_numbered(doc: Any, index: int, text: str) -> None:
    doc.add_paragraph(f"{index}. {text}")


def write_docx(parsed: ParsedKB, out_path: Path) -> None:
    doc = Document()

    doc.add_heading(parsed.title.replace(".md", "") + " - Unified Learning Material", level=0)
    doc.add_paragraph(f"Topic Focus: {parsed.topic}")
    doc.add_paragraph("Role Lens: Cross-Cloud System Architect + Cloud Application Developer")
    doc.add_paragraph("Style: Book-ready blog chapter with architecture reasoning and implementation flow.")

    doc.add_heading("How to Use This Chapter", level=1)
    doc.add_paragraph(
        "This chapter transforms the source knowledge base into an end-to-end learning narrative. "
        "Every major module, concept cluster, service, resource trail, and operational signal from the markdown source is carried forward."
    )
    doc.add_paragraph(
        "Cross-cloud parity is explicit in each section so you can think capability-first across AWS, Azure, and GCP."
    )

    doc.add_heading("Cross-Cloud Service Parity", level=1)
    if parsed.services:
        for service in parsed.services:
            azure, gcp = SERVICE_EQUIVALENTS.get(service, ("Azure equivalent varies by pattern", "GCP equivalent varies by pattern"))
            add_bullet(doc, f"{service} -> Azure: {azure} | GCP: {gcp}")
    else:
        doc.add_paragraph("No explicit service list was detected in the source markdown.")

    doc.add_heading("Module-by-Module Deep Learning", level=1)
    if parsed.modules:
        for i, module in enumerate(parsed.modules, start=1):
            doc.add_heading(f"Module {i}: {module.name}", level=2)
            doc.add_paragraph(
                "Architect View: This module should be studied as a design surface where data-path, control-path, "
                "security, scaling, and failure behavior are evaluated together."
            )

            if module.transcript_notes:
                doc.add_heading("Learning Narrative", level=3)
                for note in module.transcript_notes:
                    add_bullet(doc, f"Practical takeaway: {note}")

            if module.polished_notes:
                doc.add_heading("Architecture Clarifications", level=3)
                for note in module.polished_notes:
                    add_bullet(doc, f"Architecture note: {note}")

            if module.synthesis:
                doc.add_heading("Design Synthesis", level=3)
                for note in module.synthesis:
                    add_bullet(doc, f"System implication: {note}")

            module_services = set(parsed.services)
            if not module_services:
                guessed: list[str] = []
                for note in module.transcript_notes[:60]:
                    guessed.extend(detect_service_in_text(note))
                module_services = set(guessed)
            if module_services:
                doc.add_heading("Cross-Cloud Translation for This Module", level=3)
                for service in sorted(module_services):
                    azure, gcp = SERVICE_EQUIVALENTS.get(service, ("Azure equivalent varies by pattern", "GCP equivalent varies by pattern"))
                    add_bullet(doc, f"{service}: Azure -> {azure}; GCP -> {gcp}")

            if module.source_files:
                doc.add_heading("Resource Trail (Source Transcripts)", level=3)
                for src in module.source_files:
                    add_bullet(doc, src)
    else:
        doc.add_paragraph("No module blocks were found; source appears to be a non-topic index markdown.")

    doc.add_heading("Hands-On Execution Blueprint", level=1)
    if parsed.hands_on_activities:
        for activity in parsed.hands_on_activities:
            doc.add_heading(str(activity["title"]), level=2)
            meta_items = activity.get("meta", [])
            step_items = activity.get("steps", [])
            for meta in meta_items if isinstance(meta_items, list) else []:
                add_bullet(doc, str(meta))
            for idx, step in enumerate(step_items if isinstance(step_items, list) else [], start=1):
                add_numbered(doc, idx, f"Step {idx}: {step}")
    else:
        doc.add_paragraph("No explicit activity block was found in the source markdown.")

    doc.add_heading("Concept Coverage", level=1)
    if parsed.concepts:
        for concept in parsed.concepts:
            add_bullet(doc, concept)
    else:
        doc.add_paragraph("No explicit concept list detected.")

    doc.add_heading("Cross-Domain Resource Referrals", level=1)
    if parsed.referrals:
        for ref in parsed.referrals:
            add_bullet(doc, ref)
    else:
        doc.add_paragraph("No referral entries found.")

    doc.add_heading("Completeness Appendix - Line-by-Line Key Insight Ledger", level=1)
    doc.add_paragraph(
        "The following ledger preserves every extracted Key Insights statement from the markdown source "
        "as a study checklist so no topic signal is dropped."
    )
    if parsed.key_insights:
        for idx, insight in enumerate(parsed.key_insights, start=1):
            doc.add_paragraph(f"Insight {idx}: {insight}")
    else:
        doc.add_paragraph("No Key Insights lines were detected in the source markdown.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def generate_pack(kb_root: Path, output_root: Path) -> list[Path]:
    md_files = sorted(kb_root.glob("*.md"))
    generated: list[Path] = []

    target_root = output_root / kb_root.name
    target_root.mkdir(parents=True, exist_ok=True)

    for md_path in md_files:
        parsed = parse_md(md_path)
        out_name = md_path.stem + ".docx"
        out_path = target_root / out_name
        write_docx(parsed, out_path)
        generated.append(out_path)

    index_doc = Document()
    index_doc.add_heading(f"Learning Material Pack - {kb_root.name}", level=0)
    index_doc.add_paragraph("This index lists all generated learning DOCX chapters.")
    for out in generated:
        add_bullet(index_doc, out.name)
    index_path = target_root / "00_Learning_Pack_Index.docx"
    index_doc.save(str(index_path))
    generated.append(index_path)

    return generated


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Generate blog-style DOCX learning pack from KB markdown files.")
    p.add_argument("--kb-root", required=True, help="Path to KB folder containing markdown files.")
    p.add_argument("--output-root", default="Learning Material", help="Output root folder for DOCX learning pack.")
    return p


def main() -> None:
    args = build_parser().parse_args()
    kb_root = Path(args.kb_root)
    output_root = Path(args.output_root)

    if not kb_root.exists():
        raise SystemExit(f"KB root not found: {kb_root}")

    generated = generate_pack(kb_root=kb_root, output_root=output_root)
    print(f"Generated {len(generated)} DOCX files")
    for path in generated:
        print(f"- {path}")


if __name__ == "__main__":
    main()
